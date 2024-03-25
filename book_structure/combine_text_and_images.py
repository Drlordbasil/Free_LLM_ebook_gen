import docx
import os
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

class join_into_docx_format:
    def __init__(self, text, image_path, book_folder, chapter_title="Chapter 1", image_width=6):
        self.text = text.replace('\n', ' ').replace('\r', '')  # Remove line breaks from the text
        self.image_path = image_path
        self.book_folder = book_folder
        self.chapter_title = chapter_title
        self.image_width = image_width
        self.doc = docx.Document()

    def add_text(self):
        try:
            # Add chapter title with Heading 1 style
            heading = self.doc.add_heading(self.chapter_title, level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            heading.style = self.doc.styles.add_style('Chapter Heading', WD_STYLE_TYPE.PARAGRAPH)
            heading.style.font.name = 'Arial'
            heading.style.font.size = Pt(16)
            heading.style.font.bold = True

            # Add chapter content with justified alignment and consistent spacing
            para = self.doc.add_paragraph()
            para.add_run(self.text)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            para.style = self.doc.styles.add_style('Chapter Content', WD_STYLE_TYPE.PARAGRAPH)
            para.style.font.name = 'Times New Roman'
            para.style.font.size = Pt(12)
            para.paragraph_format.space_before = Pt(0)  # Remove space before the paragraph
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.paragraph_format.line_spacing = 1.5  # Adjust line spacing for better readability

        except Exception as e:
            print(f"Error adding text to the document: {str(e)}")
            raise

    def add_image(self):
        try:
            # Add image centered and with specified width
            para = self.doc.add_paragraph()
            run = para.add_run()
            run.add_picture(self.image_path, width=Inches(self.image_width))
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except FileNotFoundError as e:
            print(f"Error adding image to the document: Image file not found - {str(e)}")
            raise
        except Exception as e:
            print(f"Error adding image to the document: {str(e)}")
            raise

    def save_doc(self):
        try:
            output_path = os.path.join(self.book_folder, "output.docx")
            self.doc.save(output_path)
            print(f"Generated docx file: {output_path}")
        except Exception as e:
            print(f"Error saving the document: {str(e)}")
            raise